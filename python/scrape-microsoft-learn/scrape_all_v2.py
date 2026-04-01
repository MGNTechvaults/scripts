#!/usr/bin/env python3
"""Microsoft Learn scraper v2 — haalt leerpaden op als Markdown, DOCX of TXT.

Gebruik:
    python scrape_all_v2.py --format markdown       # verplichte flag
    python scrape_all_v2.py --format docx
    python scrape_all_v2.py --format txt
    python scrape_all_v2.py                         # interactieve keuze als --format ontbreekt
    python scrape_all_v2.py --config mijn.yaml      # andere config
    python scrape_all_v2.py --urls URL1 URL2        # override URLs via CLI
    python scrape_all_v2.py --output bestand        # override bestandsnaam (extensie auto-bepaald)
    python scrape_all_v2.py --headed                # browser zichtbaar
    python scrape_all_v2.py -v                      # verbose logging

Vereisten:
    pip install playwright pyyaml python-docx
    playwright install chromium
"""

import argparse
import logging
import re
import sys
import time
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from urllib.parse import urljoin

import yaml
from playwright.sync_api import sync_playwright, Page, TimeoutError as PwTimeout

logger = logging.getLogger(__name__)

VALID_FORMATS = ("markdown", "docx", "txt")
FORMAT_EXTENSIONS = {"markdown": ".md", "docx": ".docx", "txt": ".txt"}

# --- Selectors (centraal, zodat wijzigingen op één plek gebeuren) ---
CONTENT_SELECTOR = "#unit-inner-section"
FALLBACK_CONTENT_SELECTOR = "div[role='main']"
UNIT_LIST_SELECTOR = "#unit-list"
NOISE_SELECTORS = (
    "button", "footer", "nav",
    ".display-flex",
    "#next-section", "#previous-section",
    ".assessment-table",
    ".unit-page-learning-objectives",
)
MIN_CONTENT_LENGTH = 50


# ── Datamodellen ──────────────────────────────────────────────


@dataclass
class Unit:
    title: str
    url: str
    content: str


@dataclass
class Module:
    title: str
    url: str
    units: list[Unit] = field(default_factory=list)


@dataclass
class LearningPath:
    title: str
    url: str
    modules: list[Module] = field(default_factory=list)


@dataclass
class ScrapeStats:
    paths: int = 0
    modules: int = 0
    units_scraped: int = 0
    units_skipped: int = 0
    nav_failures: int = 0
    failed_urls: list[str] = field(default_factory=list)

    def summary(self) -> str:
        lines = [
            f"Paden verwerkt:    {self.paths}",
            f"Modules verwerkt:  {self.modules}",
            f"Units geschrapt:   {self.units_scraped}",
            f"Units overgeslagen: {self.units_skipped}",
            f"Navigatiefouten:   {self.nav_failures}",
        ]
        if self.failed_urls:
            lines.append("Mislukte URLs:")
            for url in self.failed_urls:
                lines.append(f"  - {url}")
        return "\n".join(lines)


# ── Hulpfuncties ──────────────────────────────────────────────


def clean_text(text: str) -> str:
    """Verwijder overbodige lege regels."""
    if not text:
        return ""
    return re.sub(r"\n\s*\n", "\n\n", text.strip())


# ── Browser helpers ───────────────────────────────────────────


def safe_goto(page: Page, url: str, *, timeout: int = 20_000, retries: int = 3) -> bool:
    """Navigeer naar URL met exponential backoff bij fouten."""
    for attempt in range(retries):
        try:
            page.goto(url, wait_until="domcontentloaded", timeout=timeout)
            return True
        except Exception as e:
            wait = 2 ** attempt
            logger.warning(
                "Navigatie mislukt (%d/%d) %s — %s, wacht %ds",
                attempt + 1, retries, url, e, wait,
            )
            time.sleep(wait)
    logger.error("Definitief mislukt: %s", url)
    return False


def wait_for_content(page: Page, timeout: int = 10_000) -> bool:
    """Wacht tot de content-selector zichtbaar is."""
    try:
        page.wait_for_selector(CONTENT_SELECTOR, state="attached", timeout=timeout)
        return True
    except PwTimeout:
        logger.debug("Content-selector niet gevonden binnen timeout op %s", page.url)
        return False


def dismiss_cookie_banner(page: Page) -> None:
    """Klik cookie-banner weg als die er is."""
    for selector in ("button#onetrust-accept-btn-handler", "button.accept-cookies"):
        try:
            page.click(selector, timeout=2000)
            logger.debug("Cookie-banner weggeklikt via %s", selector)
            return
        except PwTimeout:
            continue


def extract_content(page: Page) -> str | None:
    """Extraheert content via een DOM-clone — muteert het origineel niet."""
    noise_css = ", ".join(NOISE_SELECTORS)
    result = page.evaluate(
        """([contentSel, fallbackSel, noiseSel]) => {
            const el = document.querySelector(contentSel)
                    || document.querySelector(fallbackSel);
            if (!el) return null;
            const clone = el.cloneNode(true);
            clone.querySelectorAll(noiseSel).forEach(e => e.remove());
            return clone.innerText.trim();
        }""",
        [CONTENT_SELECTOR, FALLBACK_CONTENT_SELECTOR, noise_css],
    )
    if result and len(result) > MIN_CONTENT_LENGTH:
        return clean_text(result)
    return None


# ── URL-extractie ─────────────────────────────────────────────


def normalize_url(base: str, href: str) -> str:
    """Maak een volledige URL zonder query-params, met trailing slash gestript."""
    return urljoin(base, href).split("?")[0].rstrip("/")


def get_module_urls(page: Page) -> list[str]:
    """Haal unieke module-URLs op van een learning-path pagina."""
    links = page.query_selector_all("a[href*='/modules/']")
    seen: set[str] = set()
    urls: list[str] = []
    for link in links:
        href = link.get_attribute("href")
        if not href:
            continue
        full = normalize_url(page.url, href)
        if full not in seen:
            seen.add(full)
            urls.append(full)
    return urls


def get_unit_urls(page: Page, module_url: str) -> list[str]:
    """Haal unieke unit-URLs op van een module-pagina."""
    container = page.query_selector(UNIT_LIST_SELECTOR)
    if not container:
        logger.debug("Selector '%s' niet gevonden, fallback op <main>", UNIT_LIST_SELECTOR)
        container = page.query_selector("main")
    if not container:
        return []

    normalized_module = module_url.rstrip("/")
    seen: set[str] = set()
    urls: list[str] = []
    for link in container.query_selector_all("a"):
        href = link.get_attribute("href")
        if not href:
            continue
        full = normalize_url(page.url, href)
        # Skip de module-link zelf en externe links
        if full == normalized_module:
            continue
        if full not in seen:
            seen.add(full)
            urls.append(full)
    return urls


# ── Core scraping ─────────────────────────────────────────────


def scrape_unit(page: Page, url: str) -> Unit | None:
    """Scrape één unit (les) en retourneer een Unit of None bij falen."""
    if not safe_goto(page, url):
        return None
    wait_for_content(page)

    h1 = page.query_selector("h1")
    title = h1.inner_text().strip() if h1 else "Onbekende les"
    content = extract_content(page)

    if not content:
        logger.info("  Geen bruikbare content: %s", url)
        return None
    return Unit(title=title, url=url, content=content)


def scrape_module(page: Page, url: str, stats: ScrapeStats, seen_urls: set[str]) -> Module | None:
    """Scrape een module met al zijn units."""
    if not safe_goto(page, url):
        stats.nav_failures += 1
        stats.failed_urls.append(url)
        return None

    h1 = page.query_selector("h1")
    title = h1.inner_text().strip() if h1 else "Onbekende module"
    module = Module(title=title, url=url)

    unit_urls = get_unit_urls(page, url)
    logger.info("    %d units gevonden in '%s'", len(unit_urls), title)

    for idx, u_url in enumerate(unit_urls, 1):
        if u_url in seen_urls:
            logger.debug("    Unit al gezien, skip: %s", u_url)
            continue
        seen_urls.add(u_url)

        logger.info("      [%d/%d] %s", idx, len(unit_urls), u_url.split("/")[-1])
        unit = scrape_unit(page, u_url)
        if unit:
            module.units.append(unit)
            stats.units_scraped += 1
        else:
            stats.units_skipped += 1

    stats.modules += 1
    return module


def scrape_path(page: Page, url: str, stats: ScrapeStats, seen_urls: set[str]) -> LearningPath | None:
    """Scrape een learning path met al zijn modules en units."""
    logger.info("Leerpad: %s", url.split("/")[-2])
    if not safe_goto(page, url):
        stats.nav_failures += 1
        stats.failed_urls.append(url)
        return None

    h1 = page.query_selector("h1")
    title = h1.inner_text().strip() if h1 else "Onbekend pad"
    path = LearningPath(title=title, url=url)

    module_urls = get_module_urls(page)
    logger.info("  %d modules gevonden in '%s'", len(module_urls), title)

    for m_idx, m_url in enumerate(module_urls, 1):
        logger.info("  [%d/%d] Module verwerken...", m_idx, len(module_urls))
        module = scrape_module(page, m_url, stats, seen_urls)
        if module and module.units:
            path.modules.append(module)

    stats.paths += 1
    return path


# ── Writers ───────────────────────────────────────────────────


def _atomic_write_text(tmp: Path, output: Path, content_fn) -> None:
    """Schrijf via een tijdelijk bestand en rename atomisch bij succes."""
    with open(tmp, "w", encoding="utf-8") as f:
        content_fn(f)
    tmp.replace(output)


def write_markdown(
    paths: list[LearningPath],
    output: Path,
    title: str = "AZ-305 Cursusmateriaal",
) -> None:
    """Schrijf naar Markdown (.md). Atomisch via .tmp bestand."""
    tmp = output.with_suffix(".tmp")

    def _write(f):
        f.write(f"# {title}\n\n")
        f.write(f"Gegenereerd: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")

        f.write("## Inhoudsopgave\n\n")
        for lp in paths:
            f.write(f"- **{lp.title}**\n")
            for mod in lp.modules:
                f.write(f"  - {mod.title} ({len(mod.units)} lessen)\n")
        f.write("\n---\n\n")

        for lp in paths:
            f.write(f"## {lp.title}\n\n")
            f.write(f"Bron: {lp.url}\n\n")
            for mod in lp.modules:
                f.write(f"### {mod.title}\n\n")
                for unit in mod.units:
                    f.write(f"#### {unit.title}\n\n")
                    f.write(f"{unit.content}\n\n")
                    f.write("---\n\n")
            f.flush()

    _atomic_write_text(tmp, output, _write)
    logger.info("Markdown geschreven naar %s", output)


def write_txt(
    paths: list[LearningPath],
    output: Path,
    title: str = "AZ-305 Cursusmateriaal",
) -> None:
    """Schrijf naar platte tekst (.txt). Atomisch via .tmp bestand."""
    tmp = output.with_suffix(".tmp")
    separator = "=" * 72
    thin_sep = "-" * 72

    def _write(f):
        f.write(f"{title}\n")
        f.write(f"{separator}\n")
        f.write(f"Gegenereerd: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")

        f.write("INHOUDSOPGAVE\n")
        f.write(f"{thin_sep}\n")
        for lp in paths:
            f.write(f"  {lp.title}\n")
            for mod in lp.modules:
                f.write(f"    - {mod.title} ({len(mod.units)} lessen)\n")
        f.write(f"\n{separator}\n\n")

        for lp in paths:
            f.write(f"{lp.title.upper()}\n")
            f.write(f"{separator}\n")
            f.write(f"Bron: {lp.url}\n\n")
            for mod in lp.modules:
                f.write(f"{mod.title}\n")
                f.write(f"{thin_sep}\n\n")
                for unit in mod.units:
                    f.write(f"  {unit.title}\n")
                    f.write(f"  {'·' * (len(unit.title) + 2)}\n\n")
                    # Inspringen van content-regels
                    for line in unit.content.splitlines():
                        f.write(f"  {line}\n")
                    f.write(f"\n{thin_sep}\n\n")
            f.flush()

    _atomic_write_text(tmp, output, _write)
    logger.info("TXT geschreven naar %s", output)


def write_docx(
    paths: list[LearningPath],
    output: Path,
    title: str = "AZ-305 Cursusmateriaal",
) -> None:
    """Schrijf naar Word-document (.docx)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error(
            "python-docx is niet geïnstalleerd. Voer uit: pip install python-docx"
        )
        raise SystemExit(1)

    doc = Document()

    # Documenttitel
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datum
    date_para = doc.add_paragraph(
        f"Gegenereerd: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Inhoudsopgave (tekstueel)
    doc.add_heading("Inhoudsopgave", level=1)
    for lp in paths:
        toc_path = doc.add_paragraph(style="List Bullet")
        toc_path.add_run(lp.title).bold = True
        for mod in lp.modules:
            toc_mod = doc.add_paragraph(style="List Bullet 2")
            toc_mod.add_run(f"{mod.title} ({len(mod.units)} lessen)")

    doc.add_page_break()

    # Content
    for lp in paths:
        doc.add_heading(lp.title, level=1)
        source_para = doc.add_paragraph()
        source_para.add_run("Bron: ").bold = True
        source_para.add_run(lp.url)

        for mod in lp.modules:
            doc.add_heading(mod.title, level=2)

            for unit in mod.units:
                doc.add_heading(unit.title, level=3)
                # Splits content op lege regels — elk blok = eigen paragraaf
                for block in unit.content.split("\n\n"):
                    block = block.strip()
                    if block:
                        doc.add_paragraph(block)

                doc.add_paragraph("─" * 40)

    tmp = output.with_suffix(".tmp")
    doc.save(tmp)
    tmp.replace(output)
    logger.info("DOCX geschreven naar %s", output)


def write_output(
    paths: list[LearningPath],
    output: Path,
    fmt: str,
    title: str,
) -> None:
    """Dispatcher: kies de juiste writer op basis van het opgegeven formaat."""
    writers = {
        "markdown": write_markdown,
        "txt": write_txt,
        "docx": write_docx,
    }
    writers[fmt](paths, output, title)


# ── Config & CLI ──────────────────────────────────────────────


def load_config(path: str) -> dict:
    """Laad YAML configuratie. Retourneert lege dict als bestand niet bestaat."""
    try:
        with open(path) as f:
            return yaml.safe_load(f) or {}
    except FileNotFoundError:
        logger.debug("Config bestand '%s' niet gevonden, gebruik defaults", path)
        return {}


def ask_format_interactively() -> str:
    """Vraag de gebruiker interactief om een outputformaat als --format niet is opgegeven."""
    options = {"1": "markdown", "2": "docx", "3": "txt"}
    print()
    print("Kies een outputformaat")
    print("=" * 42)
    print("  1  ->  Markdown  (.md)")
    print("  2  ->  Word      (.docx)")
    print("  3  ->  Platte tekst (.txt)")
    print()

    while True:
        try:
            keuze = input("Voer 1, 2 of 3 in (of typ 'markdown', 'docx', 'txt'): ").strip().lower()
        except (EOFError, KeyboardInterrupt):
            print("\nAfgebroken.")
            raise SystemExit(0)

        # Accepteer zowel cijfer als naam
        if keuze in options:
            gekozen = options[keuze]
            print(f"-> Gekozen formaat: {gekozen}\n")
            return gekozen
        if keuze in VALID_FORMATS:
            print(f"-> Gekozen formaat: {keuze}\n")
            return keuze

        print(f"  Ongeldige keuze '{keuze}'. Kies 1, 2, 3 of typ het formaat.")


def resolve_format(raw: str | None) -> str:
    """Valideer het formaat. Bij None: vraag interactief. Bij ongeldig: fout + exit."""
    if raw is None:
        return ask_format_interactively()
    fmt = raw.strip().lower()
    if fmt not in VALID_FORMATS:
        print(
            f"\nFout: '{raw}' is geen geldig formaat.\n"
            f"Kies één van: {', '.join(VALID_FORMATS)}\n\n"
            f"Gebruik:  python scrape_all_v2.py --format markdown\n"
            f"          python scrape_all_v2.py --format docx\n"
            f"          python scrape_all_v2.py --format txt\n",
            file=sys.stderr,
        )
        raise SystemExit(1)
    return fmt


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description=(
            "Microsoft Learn scraper — haalt leerpaden op als Markdown, DOCX of TXT.\n\n"
            "  --format is verplicht. Als je het weglaat word je interactief om een keuze gevraagd."
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Voorbeelden:\n"
            "  python scrape_all_v2.py --format markdown\n"
            "  python scrape_all_v2.py --format docx --output MijnCursus\n"
            "  python scrape_all_v2.py --format txt --config andere.yaml -v\n"
        ),
    )
    p.add_argument(
        "--format",
        dest="fmt",
        choices=VALID_FORMATS,
        metavar="FORMAT",
        help=f"Outputformaat — verplicht. Kies uit: {', '.join(VALID_FORMATS)}",
    )
    p.add_argument("--config", default="config.yaml", help="YAML config bestand (default: config.yaml)")
    p.add_argument("--urls", nargs="+", help="Learning path URLs (overschrijft config)")
    p.add_argument("--output", help="Bestandsnaam zonder extensie, of volledig pad (extensie wordt automatisch bepaald)")
    p.add_argument("--title", help="Titel van het output document")
    p.add_argument("--headed", action="store_true", help="Browser zichtbaar maken (voor debugging)")
    p.add_argument("-v", "--verbose", action="store_true", help="Verbose logging (debug level)")
    return p.parse_args()


# ── Main ──────────────────────────────────────────────────────


def main() -> None:
    args = parse_args()

    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.StreamHandler(),
            logging.FileHandler("scraper.log", encoding="utf-8"),
        ],
    )

    # ── Formaat bepalen (verplicht — interactief als niet opgegeven) ──
    fmt = resolve_format(args.fmt)

    # ── Config laden — CLI args hebben prioriteit boven YAML ──────────
    cfg = load_config(args.config)
    urls = args.urls or cfg.get("paths", [])
    title = args.title or cfg.get("title", "AZ-305 Cursusmateriaal")

    # Output bestandsnaam bepalen — extensie volgt altijd het formaat
    raw_output = args.output or cfg.get("output", "output")
    output_path = Path(raw_output).with_suffix(FORMAT_EXTENSIONS[fmt])

    if not urls:
        logger.error("Geen URLs opgegeven via --urls of %s", args.config)
        raise SystemExit(1)

    logger.info(
        "Start scraper — formaat: %s | %d leerpaden | output: %s",
        fmt, len(urls), output_path,
    )

    stats = ScrapeStats()
    seen_urls: set[str] = set()

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=not args.headed)
        context = browser.new_context(
            user_agent=(
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            )
        )
        page = context.new_page()

        # Cookie-banner afhandelen op eerste navigatie
        if urls:
            safe_goto(page, urls[0])
            dismiss_cookie_banner(page)

        results: list[LearningPath] = []
        for url in urls:
            lp = scrape_path(page, url, stats, seen_urls)
            if lp and lp.modules:
                results.append(lp)

        browser.close()

    if results:
        write_output(results, output_path, fmt=fmt, title=title)
    else:
        logger.warning("Geen resultaten om te schrijven.")

    # Samenvatting
    logger.info("Klaar!\n%s", stats.summary())


if __name__ == "__main__":
    main()
